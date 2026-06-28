import io
import os
import urllib.parse
import uuid
from datetime import datetime as _dt
from typing import List

import aiofiles
from fastapi import APIRouter, BackgroundTasks, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from supabase import Client

from app.database import get_supabase
from app.routers.auth import get_current_user
from app.schemas.job import FileCounts, FileResponse, FileUploadResponse, JobResponse, StartJobRequest


class BatchDeleteRequest(BaseModel):
    job_ids: List[str]

router = APIRouter(prefix="/api/jobs", tags=["jobs"])

ALLOWED_EXTENSIONS = {".ppt", ".pptx", ".doc", ".docx", ".pdf"}
MAX_FILE_SIZE = 2 * 1024 * 1024 * 1024  # 2 GB
CHUNK_SIZE = 1024 * 1024  # 1 MB
UPLOAD_BASE = "uploads"
PROPOSAL_TYPES = {"qualitative", "quantitative", "presentation"}


def _build_response(job: dict, *, with_results: bool = False) -> JobResponse:
    files = job.get("review_files") or []
    counts = FileCounts(
        qualitative=sum(1 for f in files if f["proposal_type"] == "qualitative"),
        quantitative=sum(1 for f in files if f["proposal_type"] == "quantitative"),
        presentation=sum(1 for f in files if f["proposal_type"] == "presentation"),
    )
    sup = typo = blind = comp = 0
    if with_results:
        for f in files:
            for r in (f.get("review_results") or []):
                if r["category"] == "superlative":
                    sup += 1
                elif r["category"] == "typo":
                    typo += 1
                elif r["category"] in ("blind", "blind_image"):
                    blind += 1
                elif r["category"] == "competitor":
                    comp += 1
    return JobResponse(
        id=job["id"],
        status=job["status"],
        created_at=job["created_at"],
        started_at=job.get("started_at"),
        completed_at=job.get("completed_at"),
        files=[
            FileResponse(
                id=f["id"],
                proposal_type=f["proposal_type"],
                original_filename=f["original_filename"],
                file_size_bytes=f["file_size_bytes"],
                uploaded_at=f["uploaded_at"],
            )
            for f in files
        ],
        file_counts=counts,
        superlative_count=sup,
        typo_count=typo,
        blind_count=blind,
        competitor_count=comp,
    )


def _get_job(
    job_id: str,
    user_id: str,
    sb: Client,
    *,
    with_files: bool = False,
    with_results: bool = False,
) -> dict:
    if with_results:
        select_str = "*, review_files(*, review_results(*))"
    elif with_files:
        select_str = "*, review_files(*)"
    else:
        select_str = "*"
    res = sb.table("proposal_review").select(select_str).eq("id", job_id).execute()
    if not res.data:
        raise HTTPException(status_code=404, detail="кІҖнҶ  кұҙмқ„ м°ҫмқ„ мҲҳ м—ҶмҠөлӢҲлӢӨ.")
    job = res.data[0]
    if job["user_id"] != user_id:
        raise HTTPException(status_code=403, detail="м ‘к·ј к¶Ңн•ңмқҙ м—ҶмҠөлӢҲлӢӨ.")
    return job


# в”Җв”Җ кІҖнҶ  кұҙ мғқм„ұ в”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җ
@router.post("", status_code=201)
def create_job(
    current_user: dict = Depends(get_current_user),
    sb: Client = Depends(get_supabase),
):
    res = sb.table("proposal_review").insert({
        "id": str(uuid.uuid4()),
        "user_id": current_user["id"],
        "status": "draft",
    }).execute()
    return {"job_id": res.data[0]["id"]}


# в”Җв”Җ кІҖнҶ  кұҙ лӘ©лЎқ в”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җ
@router.get("", response_model=list[JobResponse])
def list_jobs(
    current_user: dict = Depends(get_current_user),
    sb: Client = Depends(get_supabase),
):
    res = sb.table("proposal_review").select(
        "*, review_files(*, review_results(*))"
    ).eq("user_id", current_user["id"]).order("created_at", desc=True).execute()
    return [_build_response(j, with_results=True) for j in (res.data or [])]


# в”Җв”Җ кІҖнҶ  кұҙ мқјкҙ„ мӮӯм ң в”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җ
@router.post("/batch-delete")
def batch_delete_jobs(
    req: BatchDeleteRequest,
    current_user: dict = Depends(get_current_user),
    sb: Client = Depends(get_supabase),
):
    """л№Ҳ draft мһЎ л“ұмқ„ н•ң лІҲмқҳ SELECT + н•ң лІҲмқҳ DELETEлЎң мқјкҙ„ мӮӯм ң.

    н”„лЎ нҠём—”л“ң loadDashboard()м—җм„ң нҢҢмқј м—ҶлҠ” мһЎмқ„ м •лҰ¬н•  л•Ң Nк°ңмқҳ к°ңлі„ мҡ”мІӯ лҢҖмӢ 
    мқҙ м—”л“ңнҸ¬мқёнҠё 1нҡҢ нҳём¶ңлЎң мІҳлҰ¬н•ңлӢӨ.
    """
    if not req.job_ids:
        return {"deleted": 0}

    # мҶҢмң к¶Ң нҷ•мқё л°Ҹ нҢҢмқј кІҪлЎң мқјкҙ„ мЎ°нҡҢ (DB 1нҡҢ)
    res = sb.table("proposal_review").select(
        "id, user_id, review_files(storage_path)"
    ).in_("id", req.job_ids).eq("user_id", current_user["id"]).execute()
    owned_jobs = res.data or []
    if not owned_jobs:
        return {"deleted": 0}

    # лЎңм»¬ нҢҢмқј м •лҰ¬
    for job in owned_jobs:
        for f in (job.get("review_files") or []):
            try:
                path = f.get("storage_path") or ""
                if path and os.path.exists(path):
                    os.remove(path)
                parent = os.path.dirname(path) if path else ""
                if parent and os.path.isdir(parent) and not os.listdir(parent):
                    os.rmdir(parent)
            except OSError:
                pass

    # DB мқјкҙ„ мӮӯм ң вҖ” CASCADEлЎң review_files, review_results мһҗлҸҷ мӮӯм ң (DB 1нҡҢ)
    owned_ids = [j["id"] for j in owned_jobs]
    try:
        sb.table("proposal_review").delete().in_("id", owned_ids).execute()
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"мқјкҙ„ мӮӯм ң мӨ‘ мҳӨлҘҳ: {str(exc)[:200]}")

    return {"deleted": len(owned_ids)}


# в”Җв”Җ кІҖнҶ  кұҙ мғҒм„ё в”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җ
@router.get("/{job_id}", response_model=JobResponse)
def get_job(
    job_id: str,
    current_user: dict = Depends(get_current_user),
    sb: Client = Depends(get_supabase),
):
    job = _get_job(job_id, current_user["id"], sb, with_files=True)
    return _build_response(job)


# в”Җв”Җ нҢҢмқј м—…лЎңл“ң в”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җ
@router.post("/{job_id}/files", status_code=201, response_model=FileUploadResponse)
async def upload_file(
    job_id: str,
    file: UploadFile = File(...),
    proposal_type: str = Form(...),
    current_user: dict = Depends(get_current_user),
    sb: Client = Depends(get_supabase),
):
    if proposal_type not in PROPOSAL_TYPES:
        raise HTTPException(status_code=400, detail="мҳ¬л°”лҘҙм§Җ м•ҠмқҖ м ңм•Ҳм„ң мў…лҘҳмһ…лӢҲлӢӨ.")

    filename = file.filename or ""
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"н—Ҳмҡ©лҗҳм§Җ м•ҠлҠ” нҢҢмқј нҳ•мӢқмһ…лӢҲлӢӨ. н—Ҳмҡ© нҳ•мӢқ: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    job = _get_job(job_id, current_user["id"], sb)
    if job["status"] != "draft":
        raise HTTPException(status_code=409, detail="мқҙлҜё кІҖнҶ к°Җ мӢңмһ‘лҗң кұҙмқҖ нҢҢмқјмқ„ мҲҳм •н•  мҲҳ м—ҶмҠөлӢҲлӢӨ.")

    safe_name = os.path.basename(filename) or f"file{ext}"
    file_id = str(uuid.uuid4())
    dir_path = os.path.join(UPLOAD_BASE, job_id, file_id)
    os.makedirs(dir_path, exist_ok=True)

    abs_base = os.path.abspath(UPLOAD_BASE)
    abs_path = os.path.abspath(os.path.join(dir_path, safe_name))
    if not abs_path.startswith(abs_base):
        raise HTTPException(status_code=400, detail="мҳ¬л°”лҘҙм§Җ м•ҠмқҖ нҢҢмқјлӘ…мһ…лӢҲлӢӨ.")

    total_size = 0
    async with aiofiles.open(abs_path, "wb") as out:
        while True:
            chunk = await file.read(CHUNK_SIZE)
            if not chunk:
                break
            total_size += len(chunk)
            if total_size > MAX_FILE_SIZE:
                await out.close()
                os.remove(abs_path)
                raise HTTPException(status_code=400, detail="нҢҢмқј нҒ¬кё°к°Җ 2GBлҘј мҙҲкіјн•©лӢҲлӢӨ.")
            await out.write(chunk)

    sb.table("review_files").insert({
        "id": file_id,
        "job_id": job_id,
        "proposal_type": proposal_type,
        "original_filename": safe_name,
        "storage_path": abs_path,
        "file_size_bytes": total_size,
        "mime_type": file.content_type or "application/octet-stream",
    }).execute()

    return FileUploadResponse(
        file_id=file_id,
        original_filename=safe_name,
        proposal_type=proposal_type,
    )


# в”Җв”Җ нҢҢмқј мӮӯм ң в”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җ
@router.delete("/{job_id}/files/{file_id}", status_code=204)
def delete_file(
    job_id: str,
    file_id: str,
    current_user: dict = Depends(get_current_user),
    sb: Client = Depends(get_supabase),
):
    job = _get_job(job_id, current_user["id"], sb)
    if job["status"] != "draft":
        raise HTTPException(status_code=409, detail="мқҙлҜё кІҖнҶ к°Җ мӢңмһ‘лҗң кұҙмқҖ нҢҢмқјмқ„ мӮӯм ңн•  мҲҳ м—ҶмҠөлӢҲлӢӨ.")

    res = sb.table("review_files").select("*").eq("id", file_id).eq("job_id", job_id).execute()
    if not res.data:
        raise HTTPException(status_code=404, detail="нҢҢмқјмқ„ м°ҫмқ„ мҲҳ м—ҶмҠөлӢҲлӢӨ.")
    rf = res.data[0]

    if os.path.exists(rf["storage_path"]):
        os.remove(rf["storage_path"])
    parent = os.path.dirname(rf["storage_path"])
    if os.path.isdir(parent) and not os.listdir(parent):
        os.rmdir(parent)

    sb.table("review_files").delete().eq("id", file_id).execute()


# в”Җв”Җ кІҖнҶ  кұҙ мӮӯм ң в”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җ
@router.post("/{job_id}/delete")
def delete_job(
    job_id: str,
    current_user: dict = Depends(get_current_user),
    sb: Client = Depends(get_supabase),
):
    # мЎҙмһ¬ м—¬л¶Җ нҷ•мқё (м—Ҷмңјл©ҙ мқҙлҜё мӮӯм ңлҗң кІғмңјлЎң к°„мЈј вҖ” л©ұл“ұм„ұ)
    res = sb.table("proposal_review").select("id, user_id, review_files(storage_path)").eq("id", job_id).execute()
    if not res.data:
        return {"deleted": True, "job_id": job_id}

    job = res.data[0]
    if job["user_id"] != current_user["id"]:
        raise HTTPException(status_code=403, detail="м ‘к·ј к¶Ңн•ңмқҙ м—ҶмҠөлӢҲлӢӨ.")

    files = job.get("review_files") or []

    # лЎңм»¬ нҢҢмқј м •лҰ¬ (Railway мһ„мӢң нҢҢмқјмӢңмҠӨн…ң вҖ” мӢӨнҢЁн•ҙлҸ„ л¬ҙмӢң)
    for f in files:
        try:
            path = f.get("storage_path") or ""
            if path and os.path.exists(path):
                os.remove(path)
            parent = os.path.dirname(path) if path else ""
            if parent and os.path.isdir(parent) and not os.listdir(parent):
                os.rmdir(parent)
        except OSError:
            pass

    # proposal_review мӮӯм ң вҶ’ CASCADEлЎң review_files, review_results мһҗлҸҷ мӮӯм ң
    try:
        sb.table("proposal_review").delete().eq("id", job_id).execute()
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"мӮӯм ң мӨ‘ мҳӨлҘҳ: {str(exc)[:200]}")

    return {"deleted": True, "job_id": job_id}


# в”Җв”Җ кІҖнҶ  мӢңмһ‘ в”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җ
@router.post("/{job_id}/start")
def start_job(
    job_id: str,
    background_tasks: BackgroundTasks,
    req: StartJobRequest,
    current_user: dict = Depends(get_current_user),
    sb: Client = Depends(get_supabase),
):
    from app.workers.review_task import run_review_sync

    job = _get_job(job_id, current_user["id"], sb, with_files=True)
    if job["status"] != "draft":
        raise HTTPException(status_code=409, detail="мқҙлҜё кІҖнҶ к°Җ мӢңмһ‘лҗҳм—ҲмҠөлӢҲлӢӨ.")
    files = job.get("review_files") or []
    if not files:
        raise HTTPException(status_code=400, detail="нҢҢмқјмқ„ лЁјм Җ м—…лЎңл“ңн•ҳм„ёмҡ”.")

    # лЎңкі  мқҙлҜём§Җ м ҖмһҘ (base64 вҶ’ нҢҢмқј)
    logo_path: str | None = None
    if req.blind_eval and req.blind_logo_b64:
        import base64 as _b64
        logo_dir = os.path.join(UPLOAD_BASE, job_id)
        os.makedirs(logo_dir, exist_ok=True)
        ext = ".png"
        if req.blind_logo_mime:
            mime_ext = {"image/jpeg": ".jpg", "image/png": ".png", "image/gif": ".gif", "image/webp": ".webp"}
            ext = mime_ext.get(req.blind_logo_mime, ".png")
        logo_path = os.path.abspath(os.path.join(logo_dir, f"logo{ext}"))
        try:
            logo_bytes = _b64.b64decode(req.blind_logo_b64)
        except Exception:
            raise HTTPException(status_code=400, detail="лЎңкі  мқҙлҜём§Җ лҚ°мқҙн„°к°Җ мҳ¬л°”лҘҙм§Җ м•ҠмҠөлӢҲлӢӨ.")
        with open(logo_path, "wb") as f:
            f.write(logo_bytes)

    sb.table("proposal_review").update({
        "status": "pending",
        "blind_eval": req.blind_eval,
        "blind_keywords": req.blind_keywords,
        "blind_logo_path": logo_path,
        "competitor_eval": req.competitor_eval,
        "competitor_keywords": req.competitor_keywords if req.competitor_eval else [],
    }).eq("id", job_id).execute()
    background_tasks.add_task(run_review_sync, job_id)
    return {"job_id": job_id, "status": "pending"}


# в”Җв”Җ кІҖнҶ  кІ°кіј мЎ°нҡҢ в”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җ
_PROPOSAL_LABELS = {
    "qualitative": "м •м„ұм ңм•Ҳм„ң",
    "quantitative": "м •лҹүм ңм•Ҳм„ң",
    "presentation": "л°ңн‘ңліё",
}
_PROPOSAL_ORDER = ["qualitative", "quantitative", "presentation"]


@router.get("/{job_id}/results")
def get_job_results(
    job_id: str,
    current_user: dict = Depends(get_current_user),
    sb: Client = Depends(get_supabase),
):
    job = _get_job(job_id, current_user["id"], sb, with_results=True)
    files = job.get("review_files") or []

    files_by_type: dict[str, list] = {}
    for f in files:
        files_by_type.setdefault(f["proposal_type"], []).append(f)

    total_superlative = 0
    total_typo = 0
    total_blind = 0
    total_competitor = 0
    files_with_issues = 0
    proposal_types = []

    for ptype in _PROPOSAL_ORDER:
        if ptype not in files_by_type:
            continue
        file_results = []
        for f in files_by_type[ptype]:
            items = sorted(
                f.get("review_results") or [],
                key=lambda r: (r["page_number"], r["category"]),
            )
            sup = sum(1 for r in items if r["category"] == "superlative")
            typo = sum(1 for r in items if r["category"] == "typo")
            blind = sum(1 for r in items if r["category"] in ("blind", "blind_image"))
            comp = sum(1 for r in items if r["category"] == "competitor")
            total_superlative += sup
            total_typo += typo
            total_blind += blind
            total_competitor += comp
            if sup + typo + blind + comp > 0:
                files_with_issues += 1
            file_results.append({
                "file_id": f["id"],
                "original_filename": f["original_filename"],
                "total_pages": f.get("total_pages"),
                "parse_error": f.get("parse_error"),
                "superlative_count": sup,
                "typo_count": typo,
                "blind_count": blind,
                "competitor_count": comp,
                "results": [
                    {
                        "id": r["id"],
                        "category": r["category"],
                        "detected_text": r["detected_text"],
                        "suggestion": r.get("suggestion"),
                        "page_number": r["page_number"],
                        "context": r.get("context"),
                    }
                    for r in items
                ],
            })
        proposal_types.append({
            "type": ptype,
            "label": _PROPOSAL_LABELS[ptype],
            "files": file_results,
        })

    return {
        "job_id": job_id,
        "status": job["status"],
        "error_message": job.get("error_message"),
        "blind_eval": job.get("blind_eval", False),
        "blind_keywords": job.get("blind_keywords") or [],
        "has_logo": bool(job.get("blind_logo_path")),
        "competitor_eval": job.get("competitor_eval", False),
        "competitor_keywords": job.get("competitor_keywords") or [],
        "summary": {
            "total_superlative": total_superlative,
            "total_typo": total_typo,
            "total_blind": total_blind,
            "total_competitor": total_competitor,
            "files_with_issues": files_with_issues,
        },
        "proposal_types": proposal_types,
    }


# в”Җв”Җ Word лӢӨмҡҙлЎңл“ң в”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җ
def _build_word_doc(job: dict) -> io.BytesIO:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    # м ңлӘ©
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("м ңм•Ҳм„ң кІҖнҶ  кІ°кіј")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(_dt.now().strftime("кІҖнҶ  мқјмӢң: %Yл…„ %mмӣ” %dмқј"))
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()

    # мҡ”м•Ҫ
    files = job.get("review_files") or []
    blind_eval = job.get("blind_eval", False)
    total_sup = total_typo = total_blind = files_with_issues = 0
    for f in files:
        results = f.get("review_results") or []
        sup = sum(1 for x in results if x["category"] == "superlative")
        typo = sum(1 for x in results if x["category"] == "typo")
        bl = sum(1 for x in results if x["category"] in ("blind", "blind_image"))
        total_sup += sup
        total_typo += typo
        total_blind += bl
        if results:
            files_with_issues += 1

    p = doc.add_paragraph()
    p.add_run("[ кІҖнҶ  мҡ”м•Ҫ ]").font.bold = True

    summary_rows = [
        ("кө¬л¶„", "кұҙмҲҳ"),
        ("н—Ҳмң„/кіјмһҘ к°ҖлҠҘ л¬ёкө¬(мөңмғҒкёү н‘ңнҳ„)", f"{total_sup}кұҙ"),
        ("мҳӨнғҖ", f"{total_typo}кұҙ"),
    ]
    if blind_eval:
        summary_rows.append(("лё”лқјмқёл“ң нҸүк°Җ: нҡҢмӮ¬мӢқлі„м •ліҙ", f"{total_blind}кұҙ"))
    summary_rows.append(("мқҙмҠҲ нҢҢмқј", f"{files_with_issues}к°ң"))

    tbl = doc.add_table(rows=len(summary_rows), cols=2)
    tbl.style = "Table Grid"
    for idx, (k, v) in enumerate(summary_rows):
        for ci, txt in enumerate((k, v)):
            cell = tbl.rows[idx].cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(txt)
            run.font.size = Pt(10)
            if idx == 0:
                run.font.bold = True
    tbl.columns[0].width = Cm(5)
    tbl.columns[1].width = Cm(3)

    doc.add_paragraph()

    # м ңм•Ҳм„ң мң нҳ•лі„
    files_by_type: dict[str, list] = {}
    for f in files:
        files_by_type.setdefault(f["proposal_type"], []).append(f)

    for ptype in _PROPOSAL_ORDER:
        if ptype not in files_by_type:
            continue

        p = doc.add_paragraph()
        r = p.add_run(f"в–  {_PROPOSAL_LABELS[ptype]}")
        r.font.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

        for f in files_by_type[ptype]:
            results = f.get("review_results") or []
            sups = [x for x in results if x["category"] == "superlative"]
            typs = [x for x in results if x["category"] == "typo"]
            blds = [x for x in results if x["category"] == "blind"]
            bld_imgs = [x for x in results if x["category"] == "blind_image"]

            p = doc.add_paragraph()
            r = p.add_run(f"в–¶ {f['original_filename']}")
            r.font.bold = True
            r.font.size = Pt(11)
            if f.get("total_pages"):
                r2 = p.add_run(f"  ({f['total_pages']}нҺҳмқҙм§Җ)")
                r2.font.size = Pt(9)
                r2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

            if f.get("parse_error"):
                r = doc.add_paragraph().add_run(f"вҡ  нҢҢмӢұ мҳӨлҘҳ: {f['parse_error']}")
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                doc.add_paragraph()
                continue

            if not results:
                r = doc.add_paragraph().add_run("кІҖм¶ңлҗң н•ӯлӘ© м—ҶмқҢ")
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
                doc.add_paragraph()
                continue

            BLUE = RGBColor(0x1D, 0x4E, 0xD8)

            def _add_result_table(items: list, col3_header: str, col3_key: str | None, default_col3: str = "кІҖнҶ  н•„мҡ”", highlight_color=None):
                tbl = doc.add_table(rows=len(items) + 1, cols=3)
                tbl.style = "Table Grid"
                for ci, h in enumerate(["нҺҳмқҙм§Җ", "кІҖм¶ң лӮҙмҡ©", col3_header]):
                    cell = tbl.rows[0].cells[ci]
                    cell.paragraphs[0].clear()
                    rr = cell.paragraphs[0].add_run(h)
                    rr.font.bold = True
                    rr.font.size = Pt(9)
                for ri, item in enumerate(items):
                    row = tbl.rows[ri + 1]
                    ctx = (item.get("context") or item.get("detected_text") or "")[:300]
                    detected = item.get("detected_text") or ""
                    col3_val = item.get(col3_key, "") if col3_key else default_col3

                    # нҺҳмқҙм§Җ лІҲнҳё м…Җ
                    c0 = row.cells[0]; c0.paragraphs[0].clear()
                    c0.paragraphs[0].add_run(f"{item['page_number']}p").font.size = Pt(9)

                    # кІҖм¶ң лӮҙмҡ© м…Җ вҖ” detected_text ліјл“ң (+ мғүмғҒ)
                    c1 = row.cells[1]; c1.paragraphs[0].clear()
                    para1 = c1.paragraphs[0]
                    lc, lt = ctx.lower(), detected.lower()
                    idx = lc.find(lt) if lt else -1
                    if idx != -1:
                        if idx > 0:
                            rr = para1.add_run(ctx[:idx]); rr.font.size = Pt(9)
                        rr = para1.add_run(ctx[idx:idx+len(detected)])
                        rr.font.size = Pt(9); rr.font.bold = True
                        if highlight_color:
                            rr.font.color.rgb = highlight_color
                        if idx + len(detected) < len(ctx):
                            rr = para1.add_run(ctx[idx+len(detected):]); rr.font.size = Pt(9)
                    else:
                        para1.add_run(ctx).font.size = Pt(9)

                    # л№„кі /мҲҳм • м ңм•Ҳ м…Җ
                    c2 = row.cells[2]; c2.paragraphs[0].clear()
                    c2.paragraphs[0].add_run(col3_val or "").font.size = Pt(9)
                tbl.columns[0].width = Cm(2)
                tbl.columns[1].width = Cm(10)
                tbl.columns[2].width = Cm(3)
                doc.add_paragraph()

            if sups:
                r = doc.add_paragraph().add_run(f"  н—Ҳмң„/кіјмһҘ к°ҖлҠҘ л¬ёкө¬(мөңмғҒкёү н‘ңнҳ„) ({len(sups)}кұҙ)")
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
                _add_result_table(sups, "л№„кі ", None, highlight_color=BLUE)

            if typs:
                r = doc.add_paragraph().add_run(f"  мҳӨнғҖ ({len(typs)}кұҙ)")
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                _add_result_table(typs, "мҲҳм • м ңм•Ҳ", "suggestion", highlight_color=BLUE)

            if blds:
                r = doc.add_paragraph().add_run(f"  лё”лқјмқёл“ң нҸүк°Җ: нҡҢмӮ¬мӢқлі„м •ліҙ(н…ҚмҠӨнҠё) ({len(blds)}кұҙ)")
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)
                _add_result_table(blds, "л№„кі ", "detected_text", highlight_color=BLUE)

            if bld_imgs:
                r = doc.add_paragraph().add_run(f"  лё”лқјмқёл“ң нҸүк°Җ: нҡҢмӮ¬мӢқлі„м •ліҙ(мқҙлҜём§Җ) ({len(bld_imgs)}кұҙ)")
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)
                _add_result_table(bld_imgs, "л№„кі ", "detected_text", highlight_color=BLUE)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@router.get("/{job_id}/download/word")
def download_word(
    job_id: str,
    current_user: dict = Depends(get_current_user),
    sb: Client = Depends(get_supabase),
):
    job = _get_job(job_id, current_user["id"], sb, with_results=True)
    buf = _build_word_doc(job)
    filename = urllib.parse.quote(f"м ңм•Ҳм„ңкІҖнҶ кІ°кіј_{job_id[:8]}.docx")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"},
    )


# в”Җв”Җ кІҖнҶ  мһ¬мӢңлҸ„ в”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җ
@router.post("/{job_id}/retry")
def retry_job(
    job_id: str,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user),
    sb: Client = Depends(get_supabase),
):
    from app.workers.review_task import run_review_sync

    job = _get_job(job_id, current_user["id"], sb, with_files=True)
    if job["status"] != "failed":
        raise HTTPException(status_code=409, detail="failed мғҒнғңмқё кІҖнҶ  кұҙл§Ң мһ¬мӢңлҸ„н•  мҲҳ мһҲмҠөлӢҲлӢӨ.")

    file_ids = [f["id"] for f in (job.get("review_files") or [])]
    if file_ids:
        sb.table("review_results").delete().in_("file_id", file_ids).execute()
        sb.table("review_files").update({"parse_error": None, "total_pages": None}).in_("id", file_ids).execute()

    sb.table("proposal_review").update({
        "status": "pending",
        "error_message": None,
        "started_at": None,
        "completed_at": None,
    }).eq("id", job_id).execute()
    background_tasks.add_task(run_review_sync, job_id)
    return {"job_id": job_id, "status": "pending"}


# в”Җв”Җ кІҖнҶ  мғҒнғң мЎ°нҡҢ в”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җв”Җ
@router.get("/{job_id}/status")
def get_job_status(
    job_id: str,
    current_user: dict = Depends(get_current_user),
    sb: Client = Depends(get_supabase),
):
    job = _get_job(job_id, current_user["id"], sb, with_files=True)
    files = job.get("review_files") or []
    return {
        "status": job["status"],
        "started_at": job.get("started_at"),
        "completed_at": job.get("completed_at"),
        "error_message": job.get("error_message"),
        "file_statuses": [
            {
                "file_id": f["id"],
                "original_filename": f["original_filename"],
                "parse_error": f.get("parse_error"),
                "total_pages": f.get("total_pages"),
            }
            for f in files
        ],
    }
